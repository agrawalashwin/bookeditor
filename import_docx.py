#!/usr/bin/env python3
import sys
import os
sys.path.insert(0, '/Users/test/Library/Python/3.9/lib/python/site-packages')
"""
Script to import DOCX files to BookEditor with proper heading recognition
"""
import requests
import json
import sys
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

API_BASE = "http://localhost:8000"

def docx_to_markdown(docx_path):
    """Convert DOCX to markdown preserving headings and structure"""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")  # Empty line
                continue
            
            # Check if it's a heading
            if paragraph.style.name.startswith('Heading'):
                heading_level = int(paragraph.style.name.split()[-1])
                markdown_content.append(f"{'#' * heading_level} {text}")
            elif paragraph.style.name == 'Title':
                markdown_content.append(f"# {text}")
            elif paragraph.style.name == 'Subtitle':
                markdown_content.append(f"## {text}")
            else:
                # Regular paragraph
                markdown_content.append(text)
        
        return "\n\n".join(markdown_content)
        
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def extract_metadata_from_docx(docx_path):
    """Extract title and author from DOCX properties"""
    try:
        doc = Document(docx_path)
        core_props = doc.core_properties
        
        title = core_props.title or os.path.splitext(os.path.basename(docx_path))[0]
        author = core_props.author or "Unknown Author"
        
        # If no metadata, try to extract from first heading
        if not core_props.title:
            for paragraph in doc.paragraphs:
                if paragraph.style.name in ['Title', 'Heading 1'] and paragraph.text.strip():
                    title = paragraph.text.strip()
                    break
        
        return title, author
        
    except Exception as e:
        print(f"Error extracting metadata: {e}")
        return os.path.splitext(os.path.basename(docx_path))[0], "Unknown Author"

def upload_manuscript(title, author, content):
    """Upload a manuscript and trigger processing"""
    
    # Create manuscript
    print(f"Creating manuscript: {title}")
    response = requests.post(f"{API_BASE}/manuscripts/", json={
        "title": title,
        "author": author,
        "content": content
    })
    
    if response.status_code != 200:
        print(f"Error creating manuscript: {response.text}")
        return None
    
    manuscript = response.json()
    manuscript_id = manuscript["id"]
    print(f"✅ Manuscript created with ID: {manuscript_id}")
    
    # Trigger chunking and embedding
    print("🔄 Processing chunks and embeddings...")
    response = requests.post(f"{API_BASE}/manuscripts/{manuscript_id}/ingest")
    
    if response.status_code == 200:
        print("✅ Chunking and embedding completed!")
        print(f"🌐 View in browser: http://localhost:3000")
        print(f"📝 Manuscript ID: {manuscript_id}")
        
        # Save manuscript ID for the web UI
        with open("current_manuscript_id.txt", "w") as f:
            f.write(manuscript_id)
        print("💾 Manuscript ID saved to current_manuscript_id.txt")
        
        return manuscript_id
    else:
        print(f"❌ Error processing embeddings: {response.text}")
        return manuscript_id

def analyze_document_structure(content):
    """Analyze the document structure and provide insights"""
    lines = content.split('\n')
    
    headings = {
        'h1': [],
        'h2': [],
        'h3': [],
        'h4': [],
        'h5': [],
        'h6': []
    }
    
    paragraphs = 0
    
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            if level <= 6:
                heading_text = line.lstrip('# ').strip()
                headings[f'h{level}'].append(heading_text)
        elif line and not line.startswith('#'):
            paragraphs += 1
    
    return headings, paragraphs

def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python import_docx.py <file.docx> [title] [author]")
        print("")
        print("Examples:")
        print("  python import_docx.py my_book.docx")
        print("  python import_docx.py my_book.docx 'Custom Title' 'Custom Author'")
        print("")
        print("The script will:")
        print("  • Extract text and preserve heading structure")
        print("  • Convert to markdown format")
        print("  • Detect chapters and sections automatically")
        print("  • Process for AI-powered editing")
        return
    
    docx_path = sys.argv[1]
    
    if not os.path.isfile(docx_path):
        print(f"❌ File not found: {docx_path}")
        return
    
    if not docx_path.lower().endswith('.docx'):
        print(f"❌ File must be a .docx file")
        return
    
    print(f"📖 Reading DOCX file: {docx_path}")
    
    # Convert DOCX to markdown
    content = docx_to_markdown(docx_path)
    if not content:
        print("❌ Failed to read DOCX file")
        return
    
    # Extract metadata
    if len(sys.argv) > 2:
        title = sys.argv[2]
        author = sys.argv[3] if len(sys.argv) > 3 else "Unknown Author"
    else:
        title, author = extract_metadata_from_docx(docx_path)
    
    print(f"📄 Title: {title}")
    print(f"✍️  Author: {author}")
    print(f"📊 Content length: {len(content)} characters")
    
    # Analyze structure
    headings, paragraphs = analyze_document_structure(content)
    print(f"📋 Structure analysis:")
    for level, heading_list in headings.items():
        if heading_list:
            print(f"   {level.upper()}: {len(heading_list)} headings")
            for i, heading in enumerate(heading_list[:3]):  # Show first 3
                print(f"      - {heading}")
            if len(heading_list) > 3:
                print(f"      ... and {len(heading_list) - 3} more")
    print(f"   Paragraphs: {paragraphs}")
    print("")
    
    # Upload and process
    manuscript_id = upload_manuscript(title, author, content)
    
    if manuscript_id:
        print("")
        print("🎉 Success! Your DOCX manuscript is ready for AI editing.")
        print("   📝 Headings and structure preserved")
        print("   🧠 Chunked and embedded for smart context")
        print("   🌐 Open http://localhost:3000 to start editing!")
        print("")
        print("💡 Tips:")
        print("   • Headings will help the AI understand document structure")
        print("   • Chapter breaks improve context-aware suggestions")
        print("   • The AI can now reference other parts of your book")

if __name__ == "__main__":
    main()
