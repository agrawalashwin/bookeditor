import io
from typing import BinaryIO
from docx import Document
from docx.shared import Inches
from sqlalchemy.orm import Session

from ..models import Manuscript


class ExportService:
    def __init__(self):
        pass
    
    def export_to_markdown(self, db: Session, manuscript_id: str) -> str:
        """Export manuscript to markdown format."""
        manuscript = db.query(Manuscript).filter(Manuscript.id == manuscript_id).first()
        if not manuscript or not manuscript.current_version:
            raise ValueError("Manuscript or current version not found")
        
        content = manuscript.current_version.content
        
        # Add metadata header
        markdown = f"""# {manuscript.title}

**Author:** {manuscript.author or 'Unknown'}  
**Version:** {manuscript.current_version.version_tag}  
**Generated:** {manuscript.current_version.created_at.strftime('%Y-%m-%d %H:%M:%S')}

---

{content}
"""
        return markdown
    
    def export_to_docx(self, db: Session, manuscript_id: str) -> BinaryIO:
        """Export manuscript to DOCX format."""
        manuscript = db.query(Manuscript).filter(Manuscript.id == manuscript_id).first()
        if not manuscript or not manuscript.current_version:
            raise ValueError("Manuscript or current version not found")
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(manuscript.title, 0)
        
        # Add metadata
        if manuscript.author:
            doc.add_paragraph(f"Author: {manuscript.author}")
        doc.add_paragraph(f"Version: {manuscript.current_version.version_tag}")
        doc.add_paragraph(f"Generated: {manuscript.current_version.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add page break
        doc.add_page_break()
        
        # Add content
        content = manuscript.current_version.content
        
        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                # Check if it's a heading (starts with #)
                if paragraph_text.strip().startswith('#'):
                    # Extract heading level and text
                    heading_level = len(paragraph_text) - len(paragraph_text.lstrip('#'))
                    heading_text = paragraph_text.lstrip('# ').strip()
                    doc.add_heading(heading_text, min(heading_level, 9))
                else:
                    doc.add_paragraph(paragraph_text.strip())
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def get_export_filename(self, db: Session, manuscript_id: str, format: str) -> str:
        """Generate appropriate filename for export."""
        manuscript = db.query(Manuscript).filter(Manuscript.id == manuscript_id).first()
        if not manuscript:
            raise ValueError("Manuscript not found")
        
        # Sanitize title for filename
        safe_title = "".join(c for c in manuscript.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')
        
        if format.lower() == 'markdown':
            return f"{safe_title}.md"
        elif format.lower() == 'docx':
            return f"{safe_title}.docx"
        else:
            raise ValueError(f"Unsupported format: {format}")
